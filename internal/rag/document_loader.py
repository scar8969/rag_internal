from pathlib import Path
from dataclasses import dataclass
from typing import Optional
import re


@dataclass
class Document:
    text: str
    source: str
    metadata: dict


SUPPORTED_EXTENSIONS = {
    ".txt", ".md", ".markdown",
    ".pdf", ".html", ".htm",
    ".docx", ".doc",
}


class DocumentLoader:
    def load(self, source: str) -> Document:
        path = Path(source)
        
        if path.exists() and path.is_file():
            return self._load_from_file(source)
        elif source.startswith("http://") or source.startswith("https://"):
            return self._load_from_url(source)
        else:
            raise ValueError(f"Cannot load source: {source}")

    def _load_from_file(self, source: str) -> Document:
        path = Path(source)
        ext = path.suffix.lower()
        
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")
        
        if ext == ".pdf":
            return self._load_pdf(source, path)
        elif ext in {".html", ".htm"}:
            return self._load_html(source, path)
        elif ext in {".docx", ".doc"}:
            return self._load_docx(source, path)
        else:
            return self._load_text(source, path)

    def _load_text(self, source: str, path: Path) -> Document:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        
        return Document(
            text=text,
            source=source,
            metadata={
                "filename": path.name,
                "extension": path.suffix,
                "file_size": path.stat().st_size,
            },
        )

    def _load_pdf(self, source: str, path: Path) -> Document:
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            text = "\n\n".join(text_parts)
            
            return Document(
                text=text,
                source=source,
                metadata={
                    "filename": path.name,
                    "extension": path.suffix,
                    "file_size": path.stat().st_size,
                    "page_count": len(pdf.pages) if 'pdf' in dir() else 0,
                },
            )
        except ImportError:
            try:
                from PyPDF2 import PdfReader
                
                text_parts = []
                reader = PdfReader(path)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                text = "\n\n".join(text_parts)
                
                return Document(
                    text=text,
                    source=source,
                    metadata={
                        "filename": path.name,
                        "extension": path.suffix,
                        "file_size": path.stat().st_size,
                        "page_count": len(reader.pages),
                    },
                )
            except ImportError:
                raise ImportError("Install pdfplumber or PyPDF2 for PDF support")

    def _load_html(self, source: str, path: Path) -> Document:
        try:
            from bs4 import BeautifulSoup
            
            with open(path, "r", encoding="utf-8") as f:
                html = f.read()
            
            soup = BeautifulSoup(html, "html.parser")
            
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text(separator="\n", strip=True)
            text = re.sub(r"\n{3,}", "\n\n", text)
            
            return Document(
                text=text,
                source=source,
                metadata={
                    "filename": path.name,
                    "extension": path.suffix,
                    "file_size": path.stat().st_size,
                },
            )
        except ImportError:
            raise ImportError("Install beautifulsoup4 for HTML support")

    def _load_docx(self, source: str, path: Path) -> Document:
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            
            return Document(
                text=text,
                source=source,
                metadata={
                    "filename": path.name,
                    "extension": path.suffix,
                    "file_size": path.stat().st_size,
                },
            )
        except ImportError:
            raise ImportError("Install python-docx for DOCX support")

    def _load_from_url(self, url: str) -> Document:
        try:
            import httpx
            from bs4 import BeautifulSoup
            
            response = httpx.get(url, timeout=30)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, "html.parser")
            
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text(separator="\n", strip=True)
            text = re.sub(r"\n{3,}", "\n\n", text)
            
            return Document(
                text=text,
                source=url,
                metadata={
                    "url": url,
                    "content_type": response.headers.get("content-type", ""),
                },
            )
        except ImportError:
            raise ImportError("Install httpx and beautifulsoup4 for URL support")

    def load_directory(
        self,
        directory: str,
        extensions: Optional[list[str]] = None,
    ) -> list[Document]:
        path = Path(directory)
        if not path.is_dir():
            raise ValueError(f"Not a directory: {directory}")
        
        docs = []
        extensions = extensions or list(SUPPORTED_EXTENSIONS)
        extensions = [ext if ext.startswith(".") else f".{ext}" for ext in extensions]
        
        for ext in extensions:
            for file_path in path.glob(f"*{ext}"):
                if file_path.is_file():
                    try:
                        doc = self._load_from_file(str(file_path))
                        docs.append(doc)
                    except Exception:
                        pass
        
        return docs
